"""Extract text from uploaded files (PDF, DOCX, TXT, HTML)."""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".html", ".htm"}


class LoaderError(ValueError):
    pass


def load_documents(path: str | Path) -> list[Document]:
    file_path = Path(path)
    if not file_path.exists():
        raise LoaderError(f"File not found: {file_path.name}")
    if file_path.stat().st_size == 0:
        raise LoaderError(f"`{file_path.name}` is empty.")

    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise LoaderError(
            f"Unsupported format `{suffix}`. Use PDF, DOCX, TXT, or HTML."
        )

    if suffix == ".pdf":
        docs = _load_pdf(file_path)
    elif suffix == ".docx":
        docs = _load_docx(file_path)
    elif suffix in {".html", ".htm"}:
        docs = _load_html(file_path)
    else:
        docs = _load_txt(file_path)

    if not docs or not any(d.page_content.strip() for d in docs):
        raise LoaderError(f"No extractable text found in `{file_path.name}`.")
    return docs


def _load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path))
    documents: list[Document] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={"filename": path.name, "page": i, "source": str(path)},
            )
        )
    return documents


def _load_docx(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"filename": path.name, "page": 1, "source": str(path)},
        )
    ]


def _load_txt(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"filename": path.name, "page": 1, "source": str(path)},
        )
    ]


def _load_html(path: Path) -> list[Document]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = " ".join(soup.get_text(separator=" ").split())
    if not text:
        return []
    return [
        Document(
            page_content=text,
            metadata={"filename": path.name, "page": 1, "source": str(path)},
        )
    ]
